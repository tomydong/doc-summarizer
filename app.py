import os
import tempfile
import datetime
from flask import Flask, request, jsonify, render_template, send_file
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
import re
import io
import pypandoc
import fitz
from flask_cors import CORS
import google.generativeai as genai

try:
    pypandoc.get_pandoc_path()
except OSError:
    pypandoc.download_pandoc()

# Khởi tạo các biến môi trường
load_dotenv()

# Cấu hình Flask app
app = Flask(__name__)
CORS(app)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # Giới hạn kích thước file 16MB
app.config['UPLOAD_FOLDER'] = tempfile.gettempdir()  # Thư mục tạm để lưu file upload

# Định nghĩa các loại file được chấp nhận
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'doc', 'rtf', 'odt', 'md'}

# Hàm kiểm tra phần mở rộng của file
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Hàm đọc nội dung từ các loại file khác nhau
def extract_text_from_file(file_path, file_ext):
    try:
        if file_ext == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif file_ext in ['docx', 'doc', 'rtf', 'odt']:
            try:
                # Sử dụng pypandoc để chuyển đổi các định dạng văn bản phổ biến sang text
                return pypandoc.convert_file(file_path, 'plain', format=file_ext)
            except Exception as e:
                if "pandoc was not found" in str(e):
                    return "Lỗi: Không tìm thấy Pandoc. Vui lòng cài đặt Pandoc và thêm vào biến môi trường PATH."
                return f"Lỗi khi đọc file {file_ext}: {str(e)}"
        elif file_ext == 'pdf':
            try:
#                import pdftotext
#                with open(file_path, "rb") as f:
#                    pdf = pdftotext.PDF(f)
#                # Trả về toàn bộ văn bản, các trang được phân tách bằng dấu xuống dòng kép
#                return "\n\n".join(pdf)
                doc = fitz.open(file_path)
                return "\n\n".join(page.get_text() for page in doc)
            except ImportError:
                return "Lỗi: Thư viện 'pdftotext' không được cài đặt. Vui lòng cài đặt bằng lệnh: pip install pdftotext"
            except Exception as e:
                print(f"Lỗi khi đọc file PDF bằng pdftotext: {str(e)}")
                return f"Lỗi khi đọc file PDF bằng pdftotext: {str(e)}"
        elif file_ext == 'md':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            return "Định dạng file không được hỗ trợ."
    except Exception as e:
        print(f"Lỗi khi đọc file: {str(e)}")
        return f"Lỗi khi đọc file: {str(e)}"

# Hàm tóm tắt tài liệu bằng Gemini API
def summarize_with_api(content, prompt):
    # Đặt mã API từ biến môi trường
    api_key = os.getenv("GEMINI_API_KEY")
    model_version = os.getenv("MODEL_VERSION", "gemini-2.0-flash-thinking-exp-01-21")
    if not api_key:
        return "Không tìm thấy API key trong biến môi trường. Vui lòng kiểm tra file .env."
    
    try:
        # Khởi tạo Gemini API
        genai.configure(api_key=api_key)
        
        # Tạo mô hình và cài đặt cấu hình
        model = genai.GenerativeModel(
            model_name=model_version,
#            generation_config={
#                "max_output_tokens": 2000,  # Tăng số lượng tokens tối đa
#                "temperature": 0.2,
#            }
        )
        
        # Tạo yêu cầu tóm tắt với system prompt
        system_prompt = "Bạn là một trợ lý AI chuyên tóm tắt tài liệu. Hãy tóm tắt nội dung tài liệu dưới dạng Markdown. Đảm bảo kết thúc mỗi điểm dấu đầu dòng bằng một dấu chấm hoặc dấu phẩy phù hợp."
        user_message = f"{prompt}\n\nĐây là nội dung cần tóm tắt:\n\n{content}"
        
        try:
            # Kết hợp system prompt và user message
            chat = model.start_chat(history=[])
            response = chat.send_message(
                f"{system_prompt}\n\n{user_message}"
            )
            
            # Kiểm tra response trước khi truy cập
            if hasattr(response, 'text') and response.text:
                return response.text
            elif hasattr(response, 'parts') and response.parts:
                if len(response.parts) > 0 and hasattr(response.parts[0], 'text'):
                    return response.parts[0].text
                else:
                    return "Không nhận được định dạng phản hồi hợp lệ."
            else:
                return "Không nhận được phản hồi có nội dung từ API. Vui lòng thử lại."
                
        except IndexError:
            # Xử lý cụ thể lỗi index out of range
            return "Lỗi xử lý phản hồi từ API. Vui lòng thử lại hoặc kiểm tra cấu hình API."
        
    except Exception as e:
        return f"Lỗi khi tóm tắt: {str(e)}"

# Hàm để xử lý việc chuyển đổi markdown sang DOCX
def markdown_to_docx(markdown_text, original_filename=None, default_font_name='Arial'):
    # Khởi tạo document mới
    doc = docx.Document()
    
    # Thiết lập style cho toàn bộ tài liệu
    style = doc.styles['Normal']
    font = style.font
    font.name = default_font_name
    font.size = Pt(11)
    
    # Thêm tiêu đề lớn cho tài liệu
    if original_filename:
        title_text = f'TÓM TẮT TÀI LIỆU: {original_filename}'
    else:
        title_text = 'TÓM TẮT TÀI LIỆU'
        
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Phân tích cú pháp markdown
    lines = markdown_text.split('\n')
    
    current_list = []
    in_list = False
    in_code_block = False
    
    for line in lines:
        # Xử lý tiêu đề
        header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if header_match:
            level = len(header_match.group(1))
            content = header_match.group(2).strip()
            doc.add_heading(content, level=level)
            continue
            
        # Xử lý đoạn văn bản thông thường
        if not line.strip():
            if in_list:
                # Xử lý danh sách đã tích lũy
                for item in current_list:
                    list_item = doc.add_paragraph(style='List Bullet')
                    list_item.add_run(item)
                current_list = []
                in_list = False
            doc.add_paragraph()
            continue
            
        # Xử lý danh sách không có thứ tự
        list_match = re.match(r'^(\s*[-*])\s+(.+)$', line)
        if list_match:
            in_list = True
            content = list_match.group(2).strip()
            current_list.append(content)
            continue
            
        # Xử lý danh sách có thứ tự
        ordered_list_match = re.match(r'^(\s*\d+\.)\s+(.+)$', line)
        if ordered_list_match:
            in_list = True
            content = ordered_list_match.group(2).strip()
            current_list.append(content)
            continue
            
        # Xử lý đoạn code
        if line.strip() == '```':
            in_code_block = not in_code_block
            if not in_code_block:  # Kết thúc code block
                doc.add_paragraph() # Thêm đoạn trống sau code block
            continue
        
        if in_code_block:
            code_para = doc.add_paragraph()
            code_run = code_para.add_run(line)
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(9)
            continue
            
        # Xử lý văn bản thông thường
        if not in_list:
            # Xử lý đậm, nghiêng, v.v.
            para = doc.add_paragraph()
            
            # Xử lý định dạng đậm và nghiêng
            bold_italic_pattern = r'(\*\*\*|\*\*|__|\*|_)(.*?)\1'
            last_end = 0
            
            for match in re.finditer(bold_italic_pattern, line):
                # Thêm văn bản thường trước phần được định dạng
                if match.start() > last_end:
                    para.add_run(line[last_end:match.start()])
                
                marker = match.group(1)
                content = match.group(2)
                
                run = para.add_run(content)
                if marker in ('***', '___'):
                    run.bold = True
                    run.italic = True
                elif marker in ('**', '__'):
                    run.bold = True
                elif marker in ('*', '_'):
                    run.italic = True
                    
                last_end = match.end()
            
            # Thêm phần còn lại của dòng
            if last_end < len(line):
                para.add_run(line[last_end:])
    
    # Xử lý danh sách còn lại nếu có
    if in_list:
        for item in current_list:
            list_item = doc.add_paragraph(style='List Bullet')
            list_item.add_run(item)
    
    # Thêm thông tin về ngày tạo tài liệu
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(f"Tạo lúc: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Lưu tài liệu vào memory stream
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    
    return docx_io

# Route cho trang chủ
@app.route('/')
def index():
    return render_template('index.html')

# Route xử lý upload file
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'Không có file nào được gửi lên'})
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'Không có file nào được chọn'})
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        file_ext = filename.rsplit('.', 1)[1].lower()
        content = extract_text_from_file(file_path, file_ext)
        
        # Xóa file sau khi đã đọc xong
        os.remove(file_path)
        
        return jsonify({'message': 'File đã được tải lên thành công', 'content': content})
    else:
        return jsonify({'error': 'Định dạng file không được hỗ trợ'})

# Route xử lý tóm tắt tài liệu
@app.route('/summarize', methods=['POST'])
def summarize():
    data = request.get_json()
    
    if not data or 'documentContent' not in data or 'prompt' not in data:
        return jsonify({'error': 'Thiếu thông tin cần thiết'})
    
    document_content = data['documentContent']
    prompt = data['prompt']
    
    if not document_content:
        return jsonify({'error': 'Nội dung tài liệu trống'})
    
    if not prompt:
        return jsonify({'error': 'Prompt trống'})
    
    # Kiểm tra xem nội dung có chứa lỗi Pandoc không
    if "Không tìm thấy Pandoc" in document_content:
        suggestion = """
Hướng dẫn cài đặt Pandoc:
1. Tải Pandoc từ https://pandoc.org/installing.html
2. Cài đặt và đảm bảo Pandoc được thêm vào PATH
3. Khởi động lại ứng dụng
Hoặc sử dụng lệnh: pip install pypandoc
"""
        return jsonify({'error': document_content, 'suggestion': suggestion})
    
    # Tóm tắt tài liệu
    summary = summarize_with_api(document_content, prompt)
    
    # Kiểm tra xem kết quả có phải là lỗi không
    if summary.startswith("Lỗi"):
        return jsonify({'error': summary})
    
    return jsonify({'summary': summary})

# Route xử lý export sang DOCX
@app.route('/export_docx', methods=['POST'])
def export_docx():
    data = request.get_json()
    
    if not data or 'summary' not in data:
        return jsonify({'error': 'Thiếu thông tin cần thiết'})
    
    markdown_text = data['summary']
    original_filename = data.get('filename', '')
    original_font_name = data.get('original_font_name') # Nhận tên font chữ từ client
    
    if not markdown_text:
        return jsonify({'error': 'Nội dung tóm tắt trống'})
    
    try:
        # Chuyển đổi markdown sang DOCX với tên file gốc
        docx_io = markdown_to_docx(markdown_text, original_filename, original_font_name)
        
        # Tạo tên file xuất
        if original_filename:
            # Đảm bảo không có phần mở rộng trong tên file
            base_filename = original_filename.split('.')[0] if '.' in original_filename else original_filename
            download_name = f"{base_filename}_summary.docx"
        else:
            download_name = "summary.docx"
        
        # Trả về file DOCX
        return send_file(
            docx_io,
            download_name=download_name,
            as_attachment=True,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': f'Lỗi khi xuất file DOCX: {str(e)}'})

# Khởi chạy app
if __name__ == '__main__':
    app.run(debug=True)