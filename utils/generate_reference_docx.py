from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def set_style_font(style, font_name="Times New Roman", font_size=11, justify=True):
    if style.base_style is not None:
        style.base_style = None

    font = style.font
    font.name = font_name
    font.size = Pt(font_size)

    if justify and hasattr(style, 'paragraph_format'):
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    rPr = style.element.get_or_add_rPr()

    for el in rPr.findall(qn('w:rFonts')):
        rPr.remove(el)

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.insert(0, rFonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size * 2)))
    rPr.append(sz)

def create_reference_docx(output_path="reference.docx", left_margin_inches=1, right_margin_inches=1):
    doc = Document()
    font_name = "Times New Roman"

    # Thiết lập lề trái và lề phải cho section đầu tiên (và thường là duy nhất)
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(left_margin_inches)
        section.right_margin = Inches(right_margin_inches)

    # Thiết lập style Normal
    set_style_font(doc.styles['Normal'], font_name, 11)

    # Thiết lập style cho Heading 1 đến Heading 9
    for i in range(1, 10):
        try:
            heading_style = doc.styles[f'Heading {i}']
            set_style_font(heading_style, font_name, 12 + (3 - min(i, 3)) * 2)
        except KeyError:
            print(f"[!] Style 'Heading {i}' không tồn tại, bỏ qua.")

    # Tạo các đoạn văn theo style
    paragraph_styles = [
        "Normal", "Body Text", "First Paragraph", "Compact", "Title", "Subtitle", "Author", "Date",
        "Abstract", "AbstractTitle", "Bibliography", "Block Text", "Footnote Block Text",
        "Source Code", "Footnote Text", "Definition Term", "Definition", "Caption", "Table Caption",
        "Image Caption", "Figure", "Captioned Figure", "TOC Heading"
    ]

    for style_name in paragraph_styles:
        try:
            paragraph = doc.add_paragraph(f"Đây là đoạn với style: {style_name}", style=style_name)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except KeyError:
            print(f"[!] Không tìm thấy style: {style_name} (bỏ qua)")

    # Heading 1 đến Heading 9
    for i in range(1, 10):
        doc.add_heading(f"Đây là Heading {i}", level=i)

    # Character styles
    char_style_demo = {
        "Default Paragraph Font": "Mặc định",
        "Body Text Char": "Body Text Char",
        "Verbatim Char": "Code mẫu `x = 42`",
        "Footnote Reference": "Tham khảo [1]",
        "Hyperlink": "https://example.com",
        "Section Number": "1.2.3"
    }

    for style, text in char_style_demo.items():
        p = doc.add_paragraph()
        run = p.add_run(f"{text} ({style})")
        try:
            run.style = style
        except KeyError:
            print(f"[!] Không tìm thấy character style: {style} (bỏ qua)")

    # Liệt kê các style bảng có sẵn
    print("\n📄 Các style bảng có trong tài liệu:")
    available_table_styles = [s.name for s in doc.styles if s.type == WD_STYLE_TYPE.TABLE]
    for s in available_table_styles:
        print("-", s)

    # Thêm bảng với style an toàn
    table_style_name = "Table Grid" if "Table Grid" in available_table_styles else None
    if table_style_name:
        table = doc.add_table(rows=2, cols=2, style=table_style_name)
    else:
        print("[!] Không tìm thấy style 'Table Grid', sử dụng mặc định.")
        table = doc.add_table(rows=2, cols=2)

    # Thêm nội dung bảng
    headers = ["Tiêu đề", "Giá trị"]
    data = ["Tên", "Pandoc"]
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(11)

    for col, text in enumerate(data):
        cell = table.cell(1, col)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(11)

    # Lưu file
    doc.save(output_path)
    print(f"\n✅ Đã tạo file reference.docx tại: {output_path}")

if __name__ == "__main__":
    create_reference_docx("static/reference.docx")
